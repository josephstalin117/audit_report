import os
import win32com.client
import docx
import re
import pandas as pd
import xml.etree.ElementTree as et


def conv_doc2docx(filename):
    """

    :param filename: doc文档的文件名
    :return: None
    """
    # 用os.path.splitext把文件名和扩展名拆开来，分别存为filename_base和filename_ext
    filename_base,filename_ext = os.path.splitext(filename)
    # 文件名加上.docx扩展名就是转换以后的文件名了。当然实际的转换在后面完成
    filename_conv = filename_base+'.docx'
    # 下面两行是检验前面提到的两个问题
    assert '~$' not in filename_base, f"~$ in filename:{filename}"
    assert filename_ext=='.doc', f'{filename} should be .doc file'

    # 下面一行用basedir目录和文件名形成word文档的相对链接
    path = os.path.join(basedir,filename)
    # 下面一行用basedir目录和文件名形成转换后word文档（docx）的相对链接
    path_conv = os.path.join(basedir,filename_conv)
    # 下面两行是把相对链接转换成绝对链接
    ori_path_abs = os.path.abspath(path)
    conv_path_abs = os.path.abspath(path_conv)
    word = win32com.client.Dispatch('Word.application')
    try:
        doc = word.Documents.Open(ori_path_abs)
        doc.SaveAs2(conv_path_abs, FileFormat=16)
        doc.Close()
        return filename_conv
    except Exception as e:
        print(f'Fail to convert:{filename}')
        print(e)


class DocProc:
    def __init__(self, note_fname, report_fname):
        self.note_doc = docx.Document(note_fname)
        self.report_doc = docx.Document(report_fname)
        
    
    def get_note_text(self):
        for para in self.note_doc.paragraphs:
            num = self.get_report_num(para.text)
            company_name = self.get_company_name(para.text)
            print(num)
        #print(len(self.doc.paragraphs))
        #print(self.doc.paragraphs[21].text)
        #print(self.doc.paragraphs[21]._element.xml)
        #print(self.doc.paragraphs[19].text)
        #print(self.doc.paragraphs[19]._element.xml)
    
    def get_report_text(self):
        count = 0
        report_num = ""
        report_time = ""
        for para in self.report_doc.paragraphs:
            num = self.get_report_num(para.text)
            time = self.get_report_time(para)
            if num:
                report_num = num
                count += 1
            if time:
                report_time = time
                count += 1

            if count == 2:
                break

        return report_num, report_time

    
    def get_report_num(self, text):
        report_num = re.search('鲁海会济审字.*$', text)
        if report_num:
            return report_num.group()
        else:
            return None
    
    def get_report_time(self, para):
        #print(para._element.xml)
        xml_data = para._element.xml
        parser = et.parse(xml_data)
        root = parser.getroot()
        print(root)
        return "test"
    
    def get_company_name(self, text):
        company_name = re.search('^.*公司$', text)
        if company_name:
            return company_name.group()
        else:
            return None
    
    def get_company_license(self, text):
        license = re.search('鲁海会济审字.*$', text)
        if license:
            return license.group()
        else:
            return None


if __name__ == '__main__':
    line_dict = {
        "license": "",
        "company_name": "",
        "report_time": "",
        "report_num": ""
    }
    basedir = 'F:/海天会计事务所/2017年度审计报告/济宁华都房地产开发有限公司/'
    report_fname = basedir + '审计报告.doc'
    note_fname = basedir + '会计报表附注.doc'
    new_report_fname = conv_doc2docx(report_fname)
    new_note_fname = conv_doc2docx(report_fname)
    data_proc = DocProc(new_note_fname, new_report_fname)
    data_proc.get_report_text()


